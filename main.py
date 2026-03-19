from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY

from flask import Flask, render_template, request, send_file
import pandas
import os.path

import zipfile
import io

app = Flask(__name__)

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html', output="test")


# Parameter list:
# - file: the uploaded Excel file
# Form params:
# - mode: the output format (pdf, docx, both)
# - paragraph_spacing: the spacing between paragraphs in the output document (in points)
# - line_spacing: the spacing between lines in the output document (e.g., 1.5 for 1.5x line spacing)
# - title: the title of the output document (optional)
# - alignment: the text alignment in the output document (left, center, right, justify)
# - page_numbers: whether to include page numbers in the output document (yes or null)
# - font_size: the font size in the output document (in points)
#
#  request.form['name']
#  request.form.get('name', 'default_value')  # Use this to avoid KeyError if the key is missing

@app.route('/upload-file', methods=['POST'])
def upload_file():
    file = request.files['file']
    if not file or file.filename == '':
        return "No file uploaded", 400
    if not file.filename.endswith('.xlsx'):
        return "Invalid file type. Supported types: .xlsx", 400

    df = pandas.read_excel(file)
    base = Path(file.filename).stem
    mode = request.form['mode']
    if mode == 'pdf':
        pdf_buffer = convert_to_pdf(base, df)
        return send_file(pdf_buffer,
                         mimetype='application/pdf',
                         as_attachment=True,
                         download_name=base + '.pdf')
    elif mode == 'docx':
        docx_buffer = convert_to_docx(base, df)
        return send_file(docx_buffer,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True,
                         download_name=base + '.docx')
    elif mode == 'both':
        pdf_buffer = convert_to_pdf(base, df)
        docx_buffer = convert_to_docx(base, df)

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(base + '.pdf', pdf_buffer.getvalue())
            zf.writestr(base + '.docx', docx_buffer.getvalue())

        zip_buffer.seek(0)

        return send_file(zip_buffer,
                         mimetype='application/zip',
                         as_attachment=True,
                         download_name=base + '.zip')
    else:
        return "Invalid mode. Supported modes: pdf, docx, both", 400

def convert_to_pdf(base, df):
    buffer = io.BytesIO()

    align_map = {
        "left": TA_LEFT,
        "center": TA_CENTER,
        "right": TA_RIGHT,
        "justify": TA_JUSTIFY,
    }

    paragraph_spacing = int(request.form['paragraph_spacing'])
    line_spacing = float(request.form['line_spacing'])
    title = request.form.get('title', '')
    alignment = align_map.get(request.form['alignment'].lower(), TA_LEFT)
    page_numbers = request.form.get('page_numbers', 'no') == 'yes'

    styles = getSampleStyleSheet()

    normal_style = ParagraphStyle(
        'Custom',
        parent=styles['Normal'],
        alignment=alignment,
        spaceBefore=paragraph_spacing,
        spaceAfter=paragraph_spacing,
        leading=styles['Normal'].fontSize * line_spacing,
    )
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []

    for i, row in df.iterrows():
        if request.form['page_numbers'] == 'yes':
            hp = f"{base} | Row {i + 1} of {len(df)}"
            story.append(Paragraph(hp, styles['Normal']))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black, spaceAfter=6))

        for col in df.columns:
            line = f"<b>{col}:</b> {row[col]}"
            story.append(Paragraph(line, normal_style))

        if i < len(df) - 1:
            story.append(PageBreak())

        story.append(PageBreak())
        doc.build(story)

    buffer.seek(0)
    return buffer

def convert_to_docx(base, df):
    document = Document()

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    title = request.form['title']

    alignment = align_map.get(request.form['alignment'].lower(), WD_ALIGN_PARAGRAPH.LEFT)
    spacing_pt = Pt(int(request.form['paragraph_spacing']))
    line_spacing = float(request.form['line_spacing'])

    if title:
        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_before = Pt(120)
        run = title_paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(24)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('Data wygenerowania: ' + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_break(WD_BREAK.PAGE)

    for i, row in df.iterrows():
        if request.form.get('page_numbers', 'no') == 'yes':
            hp = document.add_paragraph(base + ' | Row ' + str(i + 1) + ' of ' + str(len(df)) + '\n')
            insert_hr_docx(hp)
        last_p = None
        for col in df.columns:
            p = document.add_paragraph()

            p.alignment = alignment
            p.paragraph_format.space_below = spacing_pt
            p.paragraph_format.space_before = spacing_pt
            p.paragraph_format.line_spacing = line_spacing

            r = p.add_run(col + ': ')
            r.bold = True
            r.font.size = Pt(int(request.form['font_size']))
            r= p.add_run(str(row[col]))
            r.font.size = Pt(int(request.form['font_size']))

            last_p = p
        last_p.add_run().add_break(WD_BREAK.PAGE)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Source: https://github.com/python-openxml/python-docx/issues/105#issuecomment-442786431
def insert_hr_docx(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


if __name__ == '__main__':
    app.run(debug=True)


